from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from datetime import datetime
from typing import Optional
import json
import os
import uuid

from app.database.connection import get_db
from app.models.models import User
from app.services.schedule_parser import schedule_parser
from app.services.file_parser import FileParser

router = APIRouter(prefix="/api/auth", tags=["认证与课表"])


def parse_docx_schedule_with_day(file_path: str) -> dict:
    """从docx表格中解析课表，包含星期几信息
    
    表格结构：
    第0列: 节次（上午1-2, 3-4, 下午5-6, 7-8, 晚9-11）
    第1-7列: 周一~周日
    """
    from docx import Document
    
    doc = Document(file_path)
    
    # 提取学生信息
    info = {}
    full_text_parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            full_text_parts.append(p.text.strip())
    full_text = "\n".join(full_text_parts)
    
    import re
    info_patterns = {
        "grade": r'年级[：:]\s*(\S+)',
        "college": r'学院[：:]\s*(\S+)',
        "major": r'专业[：:]\s*(\S+)',
        "class": r'班级[：:]\s*(\S+)',
        "student_id": r'学号[：:]\s*(\S+)',
        "name": r'姓名[：:]\s*(\S+)',
    }
    for key, pattern in info_patterns.items():
        match = re.search(pattern, full_text)
        if match:
            info[key] = match.group(1).strip()
    
    current_week = schedule_parser.get_current_week()
    
    # 过滤网络课
    network_keywords = ['形势与政策', '网络课', '慕课']
    note_match = re.search(r'备注[：:]\s*(.+)', full_text, re.DOTALL)
    if note_match:
        notes = note_match.group(1)
        net_matches = re.finditer(r'\d+\.\s*(.+?)(?:--|\s*任务备注)', notes)
        for m in net_matches:
            network_keywords.append(m.group(1).strip().split('--')[0].strip())
    
    # 解析表格
    courses = []
    seen = set()
    
    # 节次映射：从行标签提取起始节次
    section_map = {
        '1-2': (1, 2), '1～2': (1, 2), '1~2': (1, 2),
        '3-4': (3, 4), '3～4': (3, 4), '3~4': (3, 4),
        '5-6': (5, 6), '5～6': (5, 6), '5~6': (5, 6),
        '7-8': (7, 8), '7～8': (7, 8), '7~8': (7, 8),
        '9-11': (9, 11), '9～11': (9, 11), '9~11': (9, 11),
        '晚上': (9, 11),
    }
    
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            # 跳过表头行
            if row_idx == 0:
                continue
            
            cells = row.cells
            if len(cells) < 2:
                continue
            
            # 第0列是节次标签
            section_text = cells[0].text.strip()
            start_sec, end_sec = 1, 2  # 默认
            for key, (s, e) in section_map.items():
                if key in section_text:
                    start_sec, end_sec = s, e
                    break
            
            # 第1-7列是周一到周日
            for col_idx in range(1, min(len(cells), 8)):
                day_of_week = col_idx  # 1=周一, 2=周二, ..., 7=周日
                cell_text = cells[col_idx].text.strip()
                if not cell_text:
                    continue
                
                # 解析单元格中的课程信息
                # 格式通常是：课程名 编号\n(周次) (节次) 教室 教师
                cell_lines = cell_text.split('\n')
                i = 0
                while i < len(cell_lines):
                    line = cell_lines[i].strip()
                    if not line:
                        i += 1
                        continue
                    
                    # 匹配课程名+编号
                    m = re.match(r'^(.+?)\s+(\d{3}[A-Z]?)\s*$', line)
                    if m and i + 1 < len(cell_lines):
                        cname = m.group(1).strip()
                        ccode = m.group(2).strip()
                        next_line = cell_lines[i + 1].strip()
                        dm = re.match(r'\(([^)]*?周)\)\s*\((\d+)(?:-(\d+))?节\)\s*(.+?)\s+([\u4e00-\u9fa5a-zA-Z·]+)\s*$', next_line)
                        if dm:
                            weeks_str = dm.group(1)
                            s_sec = int(dm.group(2))
                            e_sec = int(dm.group(3)) if dm.group(3) else s_sec
                            location = dm.group(4).strip()
                            teacher = dm.group(5).strip()
                            weeks = schedule_parser.parse_weeks(weeks_str)
                            
                            is_network = any(kw in cname for kw in network_keywords)
                            if not is_network:
                                key = f"{cname}_{weeks_str}_{s_sec}_{day_of_week}"
                                if key not in seen:
                                    seen.add(key)
                                    courses.append({
                                        "course_name": cname,
                                        "course_code": ccode,
                                        "teacher": teacher,
                                        "location": location,
                                        "start_section": s_sec,
                                        "end_section": e_sec,
                                        "day_of_week": day_of_week,
                                        "weeks": weeks,
                                        "weeks_str": weeks_str,
                                    })
                            i += 2
                            continue
                    
                    # 单行匹配
                    sm = re.match(
                        r'^(.+?)\s+(\d{3}[A-Z]?)\s*\(([^)]*?周)\)\s*\((\d+)(?:-(\d+))?节\)\s*(.+?)\s+([\u4e00-\u9fa5a-zA-Z·]+)\s*$',
                        line
                    )
                    if sm:
                        cname = sm.group(1).strip()
                        ccode = sm.group(2).strip()
                        weeks_str = sm.group(3)
                        s_sec = int(sm.group(4))
                        e_sec = int(sm.group(5)) if sm.group(5) else s_sec
                        location = sm.group(6).strip()
                        teacher = sm.group(7).strip()
                        weeks = schedule_parser.parse_weeks(weeks_str)
                        
                        is_network = any(kw in cname for kw in network_keywords)
                        if not is_network:
                            key = f"{cname}_{weeks_str}_{s_sec}_{day_of_week}"
                            if key not in seen:
                                seen.add(key)
                                courses.append({
                                    "course_name": cname,
                                    "course_code": ccode,
                                    "teacher": teacher,
                                    "location": location,
                                    "start_section": s_sec,
                                    "end_section": e_sec,
                                    "day_of_week": day_of_week,
                                    "weeks": weeks,
                                    "weeks_str": weeks_str,
                                })
                        i += 1
                        continue
                    
                    # 尝试只有课程名和详情（无编号）
                    detail_match = re.match(r'\(([^)]*?周)\)\s*\((\d+)(?:-(\d+))?节\)\s*(.+?)\s+([\u4e00-\u9fa5a-zA-Z·]+)\s*$', line)
                    if detail_match and i > 0:
                        # 上一行可能是课程名
                        prev_line = cell_lines[i-1].strip() if i > 0 else ""
                        cname = prev_line if prev_line else f"课程{i}"
                        weeks_str = detail_match.group(1)
                        s_sec = int(detail_match.group(2))
                        e_sec = int(detail_match.group(3)) if detail_match.group(3) else s_sec
                        location = detail_match.group(4).strip()
                        teacher = detail_match.group(5).strip()
                        weeks = schedule_parser.parse_weeks(weeks_str)
                        
                        is_network = any(kw in cname for kw in network_keywords)
                        if not is_network:
                            key = f"{cname}_{weeks_str}_{s_sec}_{day_of_week}"
                            if key not in seen:
                                seen.add(key)
                                courses.append({
                                    "course_name": cname,
                                    "course_code": "",
                                    "teacher": teacher,
                                    "location": location,
                                    "start_section": s_sec,
                                    "end_section": e_sec,
                                    "day_of_week": day_of_week,
                                    "weeks": weeks,
                                    "weeks_str": weeks_str,
                                })
                    
                    i += 1
    
    # 如果表格解析没有结果，用文本解析作为备用
    if not courses:
        text_courses = schedule_parser.parse_full_schedule(full_text)
        for c in text_courses.get("courses", []):
            c["day_of_week"] = 1  # 默认周一（无法确定）
            courses.append(c)
    
    # 当前周课程
    current_week_courses = [c for c in courses if current_week in c.get("weeks", [])]
    
    return {
        "info": info,
        "courses": courses,
        "current_week": current_week,
        "current_week_courses": current_week_courses,
        "total_courses": len(courses),
    }


@router.post("/register")
async def register_user(
    student_id: str = Form(...),
    name: str = Form(...),
    college: str = Form(default=""),
    major: str = Form(default=""),
    grade: str = Form(default=""),
    db: AsyncSession = Depends(get_db),
):
    """用户注册/登录（填写学号姓名即可）"""
    result = await db.execute(select(User).where(User.student_id == student_id))
    user = result.scalar_one_or_none()

    if not user:
        user = User(
            student_id=student_id,
            name=name,
            college=college,
            major=major,
            grade=grade,
        )
        db.add(user)
        await db.commit()
        await db.refresh(user)
    else:
        user.name = name
        user.college = college or user.college
        user.major = major or user.major
        user.grade = grade or user.grade
        user.updated_at = datetime.utcnow()
        await db.commit()

    return {
        "user_id": user.id,
        "student_id": user.student_id,
        "name": user.name,
        "college": user.college,
        "major": user.major,
        "grade": user.grade,
    }


@router.post("/upload-schedule")
async def upload_schedule(
    file: UploadFile = File(...),
    user_id: int = Form(...),
    db: AsyncSession = Depends(get_db),
):
    """上传课表文件（docx格式），解析并存储"""
    result = await db.execute(select(User).where(User.id == user_id))
    user = result.scalar_one_or_none()
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在")

    # 保存文件
    file_ext = os.path.splitext(file.filename)[1]
    file_name = f"schedule_{user_id}_{uuid.uuid4().hex[:8]}{file_ext}"
    file_path = os.path.join("uploads", file_name)
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)

    # 解析课表 - 直接从docx表格解析（包含星期几信息）
    try:
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext in ('.docx', '.doc'):
            schedule_data = parse_docx_schedule_with_day(file_path)
        else:
            # 其他格式用文本解析
            text = await FileParser.parse_file(file_path)
            schedule_data = schedule_parser.parse_full_schedule(text)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"课表解析失败: {str(e)}")

    # 将课表数据存入用户cas_token字段（复用字段存储JSON）
    user.cas_token = json.dumps(schedule_data, ensure_ascii=False)
    await db.commit()

    return {
        "message": "课表上传解析成功",
        "info": schedule_data["info"],
        "current_week": schedule_data["current_week"],
        "total_courses": schedule_data["total_courses"],
        "current_week_courses": schedule_data["current_week_courses"],
    }


@router.get("/schedule/")
async def get_schedule_no_user(
    week: Optional[int] = Query(default=None),
):
    """获取课表（未指定用户ID时返回空课表）"""
    return {
        "schedule": [],
        "current_week": schedule_parser.get_current_week(),
        "message": "请先登录后再查看课表",
    }


@router.get("/schedule/{user_id}")
async def get_schedule(
    user_id: int,
    week: Optional[int] = Query(default=None),
    db: AsyncSession = Depends(get_db),
):
    """获取课表，可指定周次（默认当前周）"""
    result = await db.execute(select(User).where(User.id == user_id))
    user = result.scalar_one_or_none()
    if not user:
        return {
            "schedule": [],
            "current_week": schedule_parser.get_current_week(),
            "message": "用户不存在，请先登录",
        }

    if not user.cas_token:
        return {
            "schedule": [],
            "current_week": schedule_parser.get_current_week(),
            "message": "尚未上传课表",
        }

    try:
        schedule_data = json.loads(user.cas_token)
    except json.JSONDecodeError:
        return {
            "schedule": [],
            "current_week": schedule_parser.get_current_week(),
            "message": "课表数据异常",
        }

    current_week = schedule_parser.get_current_week()
    target_week = week or current_week

    # 过滤指定周次的课程
    all_courses = schedule_data.get("courses", [])
    week_courses = [c for c in all_courses if target_week in c.get("weeks", [])]

    return {
        "info": schedule_data.get("info", {}),
        "schedule": week_courses,
        "all_courses": all_courses,
        "current_week": current_week,
        "target_week": target_week,
        "total_courses": len(all_courses),
    }


@router.get("/profile/{user_id}")
async def get_profile(user_id: int, db: AsyncSession = Depends(get_db)):
    """获取用户资料"""
    result = await db.execute(select(User).where(User.id == user_id))
    user = result.scalar_one_or_none()
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在")

    return {
        "id": user.id,
        "student_id": user.student_id,
        "name": user.name,
        "college": user.college,
        "major": user.major,
        "grade": user.grade,
        "avatar_url": user.avatar_url,
    }
